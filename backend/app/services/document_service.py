import asyncio
import logging
import os
import uuid
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.document import Document, DocumentChunk
from app.services.rag_engine import RAGEngine

logger = logging.getLogger(__name__)
settings = get_settings()


def extract_text_from_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


EXTRACTORS = {
    "pdf": extract_text_from_pdf,
    "docx": extract_text_from_docx,
    "doc": extract_text_from_docx,
    "txt": extract_text_from_txt,
}


def chunk_text(
    text: str,
    chunk_size: int = 500,
    overlap: int = 75,
) -> list[str]:
    """Split text into overlapping chunks by token approximation (words)."""
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap

    return chunks


class DocumentService:
    def __init__(self, db: AsyncSession, rag_engine: RAGEngine):
        self.db = db
        self.rag = rag_engine

    async def ingest_file(
        self,
        file_path: str,
        filename: str,
        file_type: str,
        category: str | None = None,
        uploaded_by: str | None = None,
        doc_id: str | None = None,
        source_group_id: str | None = None,
    ) -> Document:
        """Process a file: extract text, chunk, embed, and store."""
        file_size = os.path.getsize(file_path)

        if doc_id:
            # Use existing document record (created by upload endpoint)
            from sqlalchemy import select as sa_select
            result = await self.db.execute(
                sa_select(Document).where(Document.id == uuid.UUID(doc_id))
            )
            doc = result.scalar_one_or_none()
            if not doc:
                raise ValueError(f"Document record not found: {doc_id}")
        else:
            # Create document record (legacy/direct calls)
            doc = Document(
                filename=filename,
                file_type=file_type,
                file_size=file_size,
                category=category,
                uploaded_by=uuid.UUID(uploaded_by) if uploaded_by else None,
                status="processing",
            )
            self.db.add(doc)
            await self.db.flush()

        try:
            # Extract text
            extractor = EXTRACTORS.get(file_type)
            if not extractor:
                raise ValueError(f"Desteklenmeyen dosya tipi: {file_type}")

            text = await asyncio.to_thread(extractor, file_path)
            if not text.strip():
                raise ValueError("Dosyadan metin çıkarılamadı")

            # Chunk text
            chunks = chunk_text(text)
            if not chunks:
                raise ValueError("Metin parçalanamadı")

            # Index in Qdrant
            point_ids = await self.rag.index_chunks(
                chunks=chunks,
                document_id=str(doc.id),
                document_name=filename,
                category=category,
                source_group_id=source_group_id or (str(doc.source_group_id) if doc.source_group_id else None),
            )

            # Store chunk metadata in PostgreSQL
            for i, (chunk_text_content, point_id) in enumerate(zip(chunks, point_ids)):
                db_chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=i,
                    content=chunk_text_content,
                    qdrant_point_id=point_id,
                    token_count=len(chunk_text_content.split()),
                )
                self.db.add(db_chunk)

            doc.status = "indexed"
            doc.chunk_count = len(chunks)
            await self.db.flush()

            logger.info(
                "Document %s ingested: %d chunks", filename, len(chunks)
            )
            return doc

        except Exception as e:
            doc.status = "error"
            await self.db.flush()
            logger.error("Failed to ingest document %s: %s", filename, e)
            raise

    async def delete_document(self, document_id: str) -> None:
        """Delete a document and its vectors."""
        doc_uuid = uuid.UUID(document_id)
        result = await self.db.execute(
            select(Document).where(Document.id == doc_uuid)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            return

        # Delete vectors from Qdrant
        await self.rag.delete_document_vectors(document_id)

        # Delete from database (cascades to chunks)
        await self.db.delete(doc)
        await self.db.flush()

        logger.info("Document %s deleted", doc.filename)

    async def reindex_document(self, document_id: str) -> Document:
        """Re-index an existing document by re-processing its chunks."""
        doc_uuid = uuid.UUID(document_id)
        result = await self.db.execute(
            select(Document).where(Document.id == doc_uuid)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise ValueError("Döküman bulunamadı")

        # Delete old vectors
        await self.rag.delete_document_vectors(document_id)

        # Get existing chunks
        chunks_result = await self.db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == doc_uuid)
            .order_by(DocumentChunk.chunk_index)
        )
        chunks = chunks_result.scalars().all()

        if not chunks:
            raise ValueError("Döküman parçaları bulunamadı")

        # Re-index
        chunk_texts = [c.content for c in chunks]
        point_ids = await self.rag.index_chunks(
            chunks=chunk_texts,
            document_id=document_id,
            document_name=doc.filename,
            category=doc.category,
            source_group_id=str(doc.source_group_id) if doc.source_group_id else None,
        )

        # Update point IDs
        for chunk, point_id in zip(chunks, point_ids):
            chunk.qdrant_point_id = point_id

        doc.status = "indexed"
        await self.db.flush()

        return doc

    async def list_documents(
        self, limit: int = 50, offset: int = 0, source_group_id: str | None = None,
    ) -> tuple[list[Document], int]:
        """List all documents with pagination, optionally filtered by source group."""
        import uuid as _uuid
        from sqlalchemy import func
        from sqlalchemy.orm import selectinload

        count_query = select(func.count(Document.id))
        list_query = select(Document).options(selectinload(Document.source_group))

        if source_group_id:
            sg_uuid = _uuid.UUID(source_group_id)
            count_query = count_query.where(Document.source_group_id == sg_uuid)
            list_query = list_query.where(Document.source_group_id == sg_uuid)

        total = (await self.db.execute(count_query)).scalar()

        result = await self.db.execute(
            list_query
            .order_by(Document.created_at.desc())
            .limit(limit)
            .offset(offset)
        )
        documents = result.scalars().all()

        return documents, total
